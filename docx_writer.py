from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_margins(section):
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def save_cover_letter(content, filename, full_name, email, phone, linkedin):
    doc = Document()
    set_margins(doc.sections[0])

    # Add name (bold + centered)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(full_name)
    name_run.bold = True
    name_run.font.name = "Times New Roman"
    name_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    name_run.font.size = Pt(14)

    # Add contact info
    contact_info = f"{email} | {phone} | {linkedin}"
    contact_para = doc.add_paragraph(contact_info)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.runs[0]
    contact_run.font.name = "Times New Roman"
    contact_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    contact_run.font.size = Pt(12)

    # Add one blank line
    doc.add_paragraph("")

    # Prepare content lines
    body_lines = content.strip().split("\n")

    paragraph_buffer = []
    for line in body_lines:
        line = line.strip().strip("*")  # Remove leading/trailing asterisks
        if not line:
            continue
        paragraph_buffer.append(line)

    inserted_sincerely = False

    for i, line in enumerate(paragraph_buffer):
        # Skip LLM-generated "Sincerely" line
        if line.lower().startswith("sincerely"):
            continue

        # Add paragraph
        para = doc.add_paragraph(line)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(0)

        run = para.runs[0]
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)

        # Add space after date
        if re.match(r"^[A-Z][a-z]+ \d{1,2}, \d{4}$", line):
            doc.add_paragraph("")  # space after date

        # Add space before "Sincerely"
        if line.strip().startswith("A job opportunity at"):
            doc.add_paragraph("")  # space before closing

        # After last paragraph, insert Sincerely + Name manually
        if i == len(paragraph_buffer) - 1 and not inserted_sincerely:
            closing_para = doc.add_paragraph(f"Sincerely,\n{full_name}")
            closing_para.paragraph_format.space_after = Pt(0)
            closing_para.paragraph_format.line_spacing = 1.15
            closing_run = closing_para.runs[0]
            closing_run.font.name = "Times New Roman"
            closing_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            closing_run.font.size = Pt(12)
            inserted_sincerely = True

    doc.save(filename)